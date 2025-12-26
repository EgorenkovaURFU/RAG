import docx
from pathlib import Path
from loguru import logger



def parse_docx(file_path: str) -> list:

    """
    Docstring for parse_docx
    
    :param file_path: Description
    :type file_path: str
    :return: Description
    :rtype: list
    """

    file_path = Path(file_path)
    logger.info(f'Parsing DOXC: {file_path}')

    try:
        doc = docx.Document(file_path)
    except Exception as ex:
        logger.error(f'Error during parsing DOCX {file_path}: {ex}')
        return []
    
    paregraphs = "\n".join(
        p.text.strip()
        for p in doc.paragraphs
        if p.text and p.text.strip())

    # TODO позже реализовать 'section', пока None
    # for p in doc.paragraphs:
    #     if p.style.name.startswith("Heading"):
    #         current_section = p.text

    #     yield {
    #         "text": p.text,
    #         "section": current_section
    #     }


    if not paregraphs:
        logger.warning(f'DOCX is empty: {file_path}')
        return []
    
    return [{'text': paregraphs,
            'path': str(file_path),
            'file_type': 'docx',

            'page': None,
            'sheet': None,
            'section': None,

            'page_type': None}]

